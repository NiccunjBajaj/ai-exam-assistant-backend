from typing import List
from fastapi import APIRouter, Depends, HTTPException, Query, Body, Response
from datetime import datetime, timezone
import requests
from sqlalchemy.orm import Session
import uuid
from uuid import UUID
from pydantic import BaseModel
from sqlalchemy import func
from sqlalchemy.dialects.postgresql import insert
from docx import Document
from markdown_it import MarkdownIt
from io import BytesIO

from DB.db_models import Notes, FlashCards, StudySessions, QuizQuestion, Message, QuizAttempt, Session as ChatSession, User
from DB.deps import db_dependency, get_db
from utils.exceptions import OutOfCreditsError
from utils.credit_refil import refill_daily_credits
from utils.credits_manager import deduct_credits
from auth.deps import get_current_user
from auth.limits import enforce_flashcards_limit, enforce_notes_limit
from utils.model1 import flashcard, note, quiz
from utils.usage_tracker import track_notes_usage, track_flashcards_usage

UTILITY_COST = 10

def init_sarvam_eval(svaram):
    global sarvam_client
    sarvam_client = svaram


class NoteUpdate(BaseModel):
    content: str

class GenReq(BaseModel):
    user_input: str
    source: str
    file_name: str
    marks: int = 5
    title: str

class GenReqQuiz(BaseModel):
    user_input: str
    source: str
    file_name: str
    marks: int = 5
    title: str
    mode: str = "short"
    
class GenReq1(BaseModel):
    marks: int = 5
    title: str
    mode: str = "short"
    session_id: uuid.UUID
    
class GenReq2(BaseModel):
    marks: int = 5
    title: str
    mode: str = "short"
    note_ids: List[uuid.UUID] = []
    flashcard_ids: List[uuid.UUID] = []

class EvaluateReq(BaseModel):
    question: str
    correct_answer: str
    user_answer: str
    marks: int

class DocxRequest(BaseModel):
    content: str
    title: str

    



router = APIRouter()

@router.post("/generate-notes")
async def generate_notes(
    req: GenReq,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
    # _ = Depends(enforce_notes_limit),
):
    user = db.query(User).filter(User.id == current_user.id).first()

    # Track notes usage
    refill_daily_credits(db, user)

    notes_text = await note(req.user_input,req.marks)
    if not notes_text:
        raise HTTPException(502, "LLM failed to generate notes")
    
    success, remaining = deduct_credits(db, user, UTILITY_COST)

    if not success:
        raise OutOfCreditsError()
    
    session = StudySessions(
    user_id=current_user.id,
    title=req.title,
    file_name=req.file_name,
    file_type="notes",
    )
    db.add(session)
    db.commit()
    db.refresh(session)

    note_obj = Notes(
        id=uuid.uuid4(),
        user_id=current_user.id,
        session_id=session.id,
        content=notes_text,
        source_type= req.source,
        created_at=datetime.now(timezone.utc)
    )
    db.add(note_obj)
    db.commit()
    
    # track_notes_usage(str(current_user.id), db)

    return {
        "session_id": str(session.id),
        "notes": notes_text,
        "credits":remaining,
    }

@router.post("/generate-flashcards")
async def generate_flashcards(
    req: GenReq,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
    # _ = Depends(enforce_flashcards_limit),
):

    user = db.query(User).filter(User.id == current_user.id).first()

     # Track flashcards usage
    refill_daily_credits(db, user)

    response = await flashcard(req.user_input,req.marks)
    if not response:
        raise HTTPException(status_code=502, detail="LLM failed to generate flashcards")
    
    success, remaining = deduct_credits(db, user, UTILITY_COST)

    if not success:
        raise OutOfCreditsError()

     # robust parser
    flashcards = []
    for block in response.split("Q:")[1:]:
        if "A:" not in block:
            continue  # skip malformed card
        question, answer = block.split("A:", 1)
        flashcards.append((question.strip(), answer.strip()))
    
    if not flashcards:
        raise HTTPException(502, "No valid flashcards found in model output")

    session = StudySessions(
    user_id=current_user.id,
    title=req.title,
    file_name=req.file_name,
    file_type="flashcard",
    )
    db.add(session)
    db.commit()
    db.refresh(session)

    for q, a in flashcards:
        card = FlashCards(
            id=uuid.uuid4(),
            user_id=current_user.id,
            session_id=session.id,
            question=q,
            answer=a,
            source_type=req.source,
            created_at=datetime.now(timezone.utc)
        )
        db.add(card)
    db.commit()
    
    # track_flashcards_usage(str(current_user.id), db)
    
    return {
        "flashcards": [{"question": q, "answer": a} for q, a in flashcards],
        "credits":remaining
    }

@router.post("/generate-quiz")
async def generate_quiz(
    req: GenReqQuiz,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):

    user = db.query(User).filter(User.id == current_user.id).first()

     # Track flashcards usage
    refill_daily_credits(db, user)

    response = await quiz(req.user_input, req.marks, req.mode)
    
    if not response:
        raise HTTPException(status_code=502, detail="LLM failed to generate quiz")
        
    success, remaining = deduct_credits(db, user, UTILITY_COST)

    if not success:
        raise OutOfCreditsError()
    
    session = StudySessions(
        user_id=current_user.id,
        title=req.title,
        file_name=req.file_name,
        file_type="quiz",
    )
    db.add(session)
    db.commit()

    questions = []
    if not response:
        raise HTTPException(status_code=500, detail="Quiz generation failed")
    blocks = response.split("Q:")[1:]

    for block in blocks:
        try:
            if req.mode == "mcq":
                # Expect: Options A-D and Answer: X
                lines = block.strip().splitlines()
                question_line = lines[0].strip()
                options = {}
                answer = None

                for line in lines[1:]:
                    if line.startswith("A."):
                        options["A"] = line[2:].strip()
                    elif line.startswith("B."):
                        options["B"] = line[2:].strip()
                    elif line.startswith("C."):
                        options["C"] = line[2:].strip()
                    elif line.startswith("D."):
                        options["D"] = line[2:].strip()
                    elif "Answer:" in line:
                        answer = line.split("Answer:")[1].strip()

                db.add(QuizQuestion(
                    user_id=current_user.id,
                    session_id=session.id,
                    question=question_line,
                    type="mcq",
                    options=[options.get(k) for k in ["A", "B", "C", "D"]],
                    correct_option=options.get(answer),  # store actual text, not just "A"
                ))

            elif req.mode == "short":  # short-answer fallback
                if "A:" in block:
                    question, answer = block.split("A:", 1)
                    db.add(QuizQuestion(
                        user_id=current_user.id,
                        session_id=session.id,
                        question=question.strip(),
                        correct_answer=answer.strip(),
                        type="short",
                    ))
        except Exception as e:
            print(f"⚠️ Skipping malformed block: {e}")
    db.commit()
    return {"session_id": str(session.id), "total": len(questions),"credits":remaining}

@router.post("/generate-quiz-from-chat")
async def generate_quiz_from_chat(
    db: db_dependency,
    req: GenReq1,
    current_user: dict = Depends(get_current_user)
):

    user = db.query(User).filter(User.id == current_user.id).first()

     # Track flashcards usage
    refill_daily_credits(db, user)

    messages = (
        db.query(Message)
        .join(ChatSession, Message.session_id == ChatSession.id)
        .filter(
            Message.session_id == req.session_id,
            ChatSession.user_id == current_user.id,
        )
        .order_by(Message.timestamp.asc())
        .all()
    )

    if not messages:
        raise HTTPException(status_code=404, detail="No chat history found")
    
    context = "\n".join(
        [f"{msg.sender.upper()}: {msg.content}" for msg in messages]
    )

    study_session = StudySessions(
        id=uuid.uuid4(),
        user_id=current_user.id,
        title=req.title,
        file_name="from_chat",
        file_type="quiz",
        created_at=datetime.now(timezone.utc),
    )
    db.add(study_session)
    db.commit()

    response = await quiz(context, req.marks, req.mode)
    if not response:
        raise HTTPException(status_code=502, detail="LLM failed to generate quiz")

    success, remaining = deduct_credits(db, user, UTILITY_COST)

    if not success:
        raise OutOfCreditsError()
    
    blocks = response.split("Q:")[1:]

    for block in blocks:
        try:
            if req.mode == "mcq":
                # Expect: Options A-D and Answer: X
                lines = block.strip().splitlines()
                question_line = lines[0].strip()
                options = {}
                answer = None

                for line in lines[1:]:
                    if line.startswith("A."):
                        options["A"] = line[2:].strip()
                    elif line.startswith("B."):
                        options["B"] = line[2:].strip()
                    elif line.startswith("C."):
                        options["C"] = line[2:].strip()
                    elif line.startswith("D."):
                        options["D"] = line[2:].strip()
                    elif "Answer:" in line:
                        answer = line.split("Answer:")[1].strip()

                db.add(QuizQuestion(
                    user_id=current_user.id,
                    session_id=study_session.id,
                    question=question_line,
                    type="mcq",
                    options=[options.get(k) for k in ["A", "B", "C", "D"]],
                    correct_option=options.get(answer),  # store actual text, not just "A"
                ))

            elif req.mode == "short":  # short-answer fallback
                if "A:" in block:
                    question, answer = block.split("A:", 1)
                    db.add(QuizQuestion(
                        user_id=current_user.id,
                        session_id=study_session.id,
                        question=question.strip(),
                        correct_answer=answer.strip(),
                        type="short",
                    ))

        except Exception as e:
            print(f"⚠️ Skipping malformed block: {e}")

    db.commit()
    return {"session_id": str(study_session.id),"credits":remaining}

@router.post("/generate-quiz-from-study")
async def generate_quiz_from_study(
    db: db_dependency,
    req: GenReq2,
    current_user: dict = Depends(get_current_user) 
):

    user = db.query(User).filter(User.id == current_user.id).first()

     # Track flashcards usage
    refill_daily_credits(db, user)
    
    # Fetch notes and flashcards for this user (only if ids provided)
    notes = (
        db.query(Notes)
        .filter(Notes.user_id == current_user.id, Notes.id.in_(req.note_ids))
        .all()
        if (req.note_ids and len(req.note_ids) > 0) else []
    )

    flashcards = (
        db.query(FlashCards)
        .filter(FlashCards.user_id == current_user.id, FlashCards.id.in_(req.flashcard_ids))
        .all()
        if (req.flashcard_ids and len(req.flashcard_ids) > 0) else []
    )

    # Fallback: if none found but IDs provided, treat provided IDs as session IDs
    if not notes and (req.note_ids and len(req.note_ids) > 0):
        notes = (
            db.query(Notes)
            .filter(Notes.user_id == current_user.id, Notes.session_id.in_(req.note_ids))
            .all()
        )

    if not flashcards and (req.flashcard_ids and len(req.flashcard_ids) > 0):
        flashcards = (
            db.query(FlashCards)
            .filter(FlashCards.user_id == current_user.id, FlashCards.session_id.in_(req.flashcard_ids))
            .all()
        )

    if not notes and not flashcards:
        raise HTTPException(status_code=400, detail="No notes or flashcards found for the provided IDs")

    # Build a rich context with separators to help the LLM
    context_parts = []
    if notes:
        context_parts.append("\n\n".join([n.content for n in notes]))
    if flashcards:
        context_parts.append("\n\n".join([f"Q: {f.question}\nA: {f.answer}" for f in flashcards]))
    context = "\n\n".join(context_parts)

    session = StudySessions(
        id=uuid.uuid4(),
        user_id=current_user.id,
        title=req.title,
        file_type="quiz",
        file_name="from_study",
        created_at=datetime.now(timezone.utc),
    )
    db.add(session)
    db.commit()

    response = await quiz(context, req.marks, req.mode)

    if not response:
        raise HTTPException(status_code=500, detail="No quiz generated")
    
    success, remaining = deduct_credits(db, user, UTILITY_COST)

    if not success:
        raise OutOfCreditsError()

    blocks = response.split("Q:")[1:]

    for block in blocks:
        try:
            if req.mode == "mcq":
                # Expect: Options A-D and Answer: X
                lines = block.strip().splitlines()
                question_line = lines[0].strip()
                options = {}
                answer = None

                for line in lines[1:]:
                    if line.startswith("A."):
                        options["A"] = line[2:].strip()
                    elif line.startswith("B."):
                        options["B"] = line[2:].strip()
                    elif line.startswith("C."):
                        options["C"] = line[2:].strip()
                    elif line.startswith("D."):
                        options["D"] = line[2:].strip()
                    elif "Answer:" in line:
                        answer = line.split("Answer:")[1].strip()

                db.add(QuizQuestion(
                    user_id=current_user.id,
                    session_id=session.id,
                    question=question_line,
                    type="mcq",
                    options=[options.get(k) for k in ["A", "B", "C", "D"]],
                    correct_option=options.get(answer),  # store actual text, not just "A"
                ))

            elif req.mode == "short":  # short-answer fallback
                if "A:" in block:
                    question, answer = block.split("A:", 1)
                    db.add(QuizQuestion(
                        user_id=current_user.id,
                        session_id=session.id,
                        question=question.strip(),
                        correct_answer=answer.strip(),
                        type="short",
                    ))

        except Exception as e:
            print(f"⚠️ Skipping malformed block: {e}")

    db.commit()
    return {"session_id": str(session.id)}


@router.get("/notes")
async def get_notes(
    db: db_dependency,
    current_user: dict = Depends(get_current_user)
):
    notes = (
        db.query(Notes)
        .filter_by(user_id=current_user.id)
        .order_by(Notes.created_at.desc())
        .all()
    )
    return [{"id": str(n.id), "content": n.content, "created_at": n.created_at.isoformat()} for n in notes]

@router.get("/flashcards")
async def get_flashcards(
    db: db_dependency,
    current_user: dict = Depends(get_current_user)
):
    cards = (
        db.query(FlashCards)
        .filter_by(user_id=current_user.id)
        .order_by(FlashCards.created_at.desc())
        .all()
    )
    return [
        {
            "id": str(c.id),
            "question": c.question,
            "answer": c.answer,
            "created_at": c.created_at.isoformat()
        }
        for c in cards
    ]

@router.get("/study-sessions",response_model=List[dict])
async def get_study_sessions(
    db : db_dependency,
    type: str = Query(None),
    current_user: dict = Depends(get_current_user),
):
    query = db.query(StudySessions).filter(StudySessions.user_id == current_user.id)

    if type:
        query = query.filter(
            (StudySessions.file_type == type) | (StudySessions.file_type == "both")
        )

    sessions = query.order_by(StudySessions.created_at.desc()).all()
    return [
        {
            "id": str(s.id),
            "title": s.title,
            "file_name": s.file_name,
            "created_at": s.created_at.isoformat(),
        } 
        for s in sessions
    ]

@router.get("/flashcards/by-session/{session_id}")
async def get_flashcards_by_session(session_id: str, db: db_dependency, current_user: dict = Depends(get_current_user)):
    cards = (
        db.query(FlashCards)
        .filter(FlashCards.session_id == session_id)
        .filter(FlashCards.user_id == current_user.id)
        .all()
    )
    return [
        {
            "id": str(c.id), 
            "question": c.question, 
            "answer": c.answer,
            "studysession": {   # 👈 match frontend expectation
                "title": c.session.title if c.session else None
            }
            } 
        for c in cards
        ]

@router.get("/notes/by-session/{session_id}")
async def get_notes_by_session(session_id: str, db: db_dependency, current_user: dict = Depends(get_current_user)):
    notes = (
        db.query(Notes)
        .filter(Notes.session_id == session_id)
        .filter(Notes.user_id == current_user.id)
        .all()
    )
    return [
        {
            "id": str(n.id),
            "session_id": str(n.session_id),
            "content": n.content,
            "studysession": {   # 👈 match frontend expectation
                "title": n.session.title if n.session else None
            }
        }
        for n in notes
    ]

@router.put("/notes/{note_id}")
def update_note(note_id: str, note_update: NoteUpdate, db: db_dependency,current_user=Depends(get_current_user)):
    note = db.query(Notes).filter(Notes.id == note_id, Notes.user_id == current_user.id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    note.content = note_update.content
    db.commit()
    db.refresh(note)
    return {"message": "Note updated successfully", "note": note}

@router.get("/quizzes/by-session/{session_id}")
async def get_quiz(session_id: str, db: db_dependency, current_user: dict = Depends(get_current_user)):
    if not session_id or session_id == "undefined":
        raise HTTPException(status_code=400, detail="Session ID is required")
    
    try:
        quiz_uuid = uuid.UUID(session_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid session ID format")

    quizzes = (
        db.query(QuizQuestion)
        .filter(QuizQuestion.session_id == quiz_uuid)
        .filter(QuizQuestion.user_id == current_user.id)
        .all()
    )
    return [
        {
            "id": str(q.id),
            "question": q.question,
            "type": q.type,
            "options": q.options,  # ✅ make sure this is sent
            "correct_option": q.correct_option,
            "correct_answer": q.correct_answer,
            "studysession": {   # 👈 match frontend expectation
                "title": q.session.title if q.session else None
            }
        }
        for q in quizzes
    ]




#Evaluator
@router.post("/evaluate-answer")
async def evaluate_quiz(
    req: EvaluateReq,
):
    system_prompt = f"""
You are **Proff-Eval**, a fair, supportive, and encouraging **Academic Evaluator**.
Your goal is to help students improve by giving kind, constructive feedback based on their answer.

====================================================
## 🎯 Your Evaluation Task
Evaluate the student’s answer using:
- the official correct answer, and  
- the mark value for the question.

Your evaluation must be **fair, forgiving, and focused on conceptual understanding**, not wording.

The quiz which you will be evaluating is of {req.marks}{"Marks" if req.marks <=10 else "Words"}
====================================================
## 🧠 Mark-Based Expectations (Flexible)
Use the mark value to decide how much detail the student *should* ideally provide:

- **2 Marks:**  
  - 1–2 lines expected  
  - Focus on core definition / fact  
  - Do NOT punish missing examples or details  

- **5 Marks:**  
  - Short paragraph expected  
  - Main idea + few supporting points  
  - Minor missing details = still “Partially Correct”  

- **10 Marks:**  
  - Detailed explanation with reasoning  
  - Examples or depth expected  
  - Missing 1–2 points shouldn’t make it “Incorrect”  
  - As long as the core concept is correct → give “Partially Correct”  

####################################################
## ✔️ Evaluation Criteria (Soft and Fair)

### 1. Correctness
- Is the main concept right?  
- Even if phrased differently or less perfectly → **Accept it**.

### 2. Completeness
- Did they cover enough points for the mark value?  
- Missing minor points → “Partially Correct”, NOT "Incorrect".  

### 3. Clarity
- Is it understandable?  
- Ignore grammar mistakes, typos, informal tone.

####################################################
## ❤️ TONALITY (VERY IMPORTANT)

Be:
- encouraging  
- respectful  
- positive  
- motivating  

Never discourage the student.  
Even if the answer is incorrect, gently explain what was missing and appreciate the effort.

####################################################
## 🔒 Output Format (STRICT)
Respond ONLY in the following structure:

Verdict: [Correct / Partially Correct / Incorrect]

Explanation: [1–3 sentences. Encourage the student.  
If correct → praise briefly.  
If partially correct → explain what was right + what can be improved.  
If incorrect → explain the correct idea gently.]

Correct Answer: [give the correct answer]

Do NOT add extra text, headings, emojis, or markdown outside this format.
"""
    user_int = " "
    if req.user_answer.strip() == "":
        user_int = "User did not answer"
    else:
        user_int = req.user_answer.strip()

    print(user_int)

    user_ans = (
        f"Question: {req.question}\n"
        f"User's Answer: {user_int}\n"
        f"Correct Answer: {req.correct_answer}"
    )

    try:
        response_raw  = sarvam_client.chat.completions(
            messages = [
                {
                    "content":system_prompt,
                    "role":"system"
                },
                {
                    "content":user_ans,
                    "role":"user"
                }
            ]
        )
                
        reply = response_raw.choices[0].message.content

        verdict = ""
        # ---------- VERDICT ----------
        if "Verdict: Correct" in reply:
            verdict = "Correct"
        elif "Verdict: Partially Correct" in reply:
            verdict = "Partially Correct"
        elif "Verdict: Incorrect" in reply:
            verdict = "Incorrect"
        else:
            verdict = "Incorrect"

        # ---------- EXPLANATION ----------
        explanation = ""
        if "Explanation:" in reply:
            explanation = reply.split("Explanation:")[1].split("Correct Answer:")[0].strip()

        # ---------- CORRECT ANSWER ----------
        correct_ans = ""
        if "Correct Answer:" in reply:
            correct_ans = reply.split("Correct Answer:")[1].strip()

        return {
            "verdict": verdict,
            "explanation": explanation,
            "correct_ans": correct_ans,
        }

    except Exception as e:
        print("SarvamAI Eval Error:", str(e))
        return {"verdict": "Unknown", "explanation": "Evaluation failed. Try again."}


@router.post("/quiz-attempts/save")
async def save_quiz_attempts(
    db: db_dependency,
    attempts: List[dict] = Body(...),
    current_user: dict = Depends(get_current_user),
):
    if not attempts:
        raise HTTPException(status_code=400, detail="No attempts provided")
    session_id = attempts[0]["session_id"]

    total_questions = db.query(QuizQuestion).filter_by(session_id=session_id).count()
    incoming_q_ids = {a["question_id"] for a in attempts}

    # Latest attempt_number for this user + session
    prev_attempt = (
        db.query(func.max(QuizAttempt.attempt_number))
        .filter_by(user_id=current_user.id, session_id=session_id)
        .scalar()
    ) or 0

    # Normalize both sets to strings
    incoming_q_ids = {str(a["question_id"]) for a in attempts}

    quiz_ids = {
        str(q.id) for q in db.query(QuizQuestion).filter_by(session_id=session_id).all()
    }

    # Compare as pure string sets
    is_full_quiz = incoming_q_ids == quiz_ids

    attempt_number = prev_attempt + 1 if is_full_quiz else prev_attempt

    for a in attempts:
        stmt = (
            insert(QuizAttempt)
            .values(
                user_id=current_user.id,
                session_id=session_id,
                question_id=a["question_id"],
                user_answer=a["user_answer"],
                verdict=a["verdict"],
                explanation=a.get("explanation", ""),
                score=1 if a["verdict"] == "Correct" else 0,
                attempt_number=attempt_number,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "question_id", "attempt_number"],
                set_={
                    "user_answer": a["user_answer"],
                    "verdict": a["verdict"],
                    "explanation": a.get("explanation", ""),
                    "score": 1 if a["verdict"] == "Correct" else 0,
                },
            )
        )
        db.execute(stmt)

    db.commit()
    correct_count = sum(1 for a in attempts if a["verdict"].strip().lower() == "correct")
    total_questions = len(attempts)
    score_percent = round((correct_count / total_questions) * 100, 2)
    return {
        "status": "saved",
        "attempt_number": attempt_number,
        "score": score_percent,
        "correct": correct_count,
        "total": total_questions,
    }

@router.get("/quiz-attempts/summary")
def get_quiz_attempts_summary(
    db: db_dependency,
    current_user: dict = Depends(get_current_user),
):

    rows = (
        db.query(
            QuizAttempt.session_id,
            func.min(StudySessions.title).label("title"),
            func.count().label("total_attempts"),
            func.max(QuizAttempt.attempt_number).label("latest_attempt"),
            func.min(QuizAttempt.created_at).label("created_at"),
        )
        .join(StudySessions, StudySessions.id == QuizAttempt.session_id)
        .filter(QuizAttempt.user_id == current_user.id)
        .group_by(QuizAttempt.session_id)
        .order_by(func.min(QuizAttempt.created_at).desc())
        .all()
    )

    return [
        {
            "session_id": r.session_id,
            "title": r.title,
            "total_attempts": r.total_attempts,
            "latest_attempt": r.latest_attempt,
            "date": r.created_at.isoformat(),
        }
        for r in rows
    ]

@router.get("/quiz-attempts/by-session/{session_id}")
def get_all_attempts_by_session(
    session_id: UUID,
    db: db_dependency,
    current_user: dict = Depends(get_current_user),
):
    attempts = (
        db.query(QuizAttempt)
        .filter_by(user_id = current_user.id, session_id = session_id)
        .order_by(QuizAttempt.attempt_number, QuizAttempt.question_id)
        .all()
    )

    return [
        {
            "question_id": a.question_id,
            "user_answer": a.user_answer,
            "verdict": a.verdict,
            "explanation": a.explanation,
            "attempt_number": a.attempt_number,
            "created_at": a.created_at.isoformat()           
        }
        for a in attempts
    ]

@router.get("/quiz/{session_id}/draft-status")
def quiz_draft_status(
    session_id: UUID,
    db: db_dependency,
    current_user: dict = Depends(get_current_user)
):
    latest_attempt = (
        db.query(func.max(QuizAttempt.attempt_number))
        .filter(
            QuizAttempt.user_id == current_user.id,
            QuizAttempt.session_id == session_id
        )
        .scalar()

    )
    if latest_attempt is None:
        latest_attempt = 0
    incorrect_count = (
        db.query(QuizAttempt)
        .filter(
            QuizAttempt.user_id == current_user.id,
            QuizAttempt.session_id == session_id,
            QuizAttempt.attempt_number == latest_attempt,
            func.trim(func.lower(QuizAttempt.verdict)) == "incorrect"
        )
        .count()
    )

    return {
        "latest_attempt_number": latest_attempt,
        "incorrect": incorrect_count,
        "can_retry_partial": incorrect_count <= 3
    }


@router.get("/quiz/{session_id}/attempt/{attempt_number}/summary")
def get_quiz_attempt_summary(
    session_id: UUID,
    attempt_number: int,
    db: db_dependency,
    current_user: dict = Depends(get_current_user),
):
    attempts = db.query(QuizAttempt).filter_by(
        session_id=session_id,
        user_id=current_user.id,
        attempt_number=attempt_number
    ).all()

    if not attempts:
        raise HTTPException(status_code=404, detail="Attempt not found")

    total = len(attempts)
    correct = sum(1 for a in attempts if a.verdict.strip().lower() == "correct")
    incorrect = total - correct
    score = round((correct / total) * 100, 2)

    return {
        "session_id": str(session_id),
        "attempt_number": attempt_number,
        "total_questions": total,
        "correct": correct,
        "incorrect": incorrect,
        "score_percent": score,
    }


@router.get("/quiz-attempts/{session_id}")
def get_quiz_attempt(
    session_id: UUID,
    db: db_dependency,
    attempt: int = Query(None, description="Attempt number (optional)"),    
    current_user: dict = Depends(get_current_user),
):
    """
    • If `attempt` is supplied → return only that attempt’s rows  
    • If `attempt` is None     → return *all* attempts for the session
    """
    q = (
        db.query(QuizAttempt)
        .filter(
            QuizAttempt.user_id == current_user.id,
            QuizAttempt.session_id == session_id,
        )
    )

    if attempt is not None:
        q = q.filter(QuizAttempt.attempt_number == attempt)

    # Join to quiz table just to fetch question text in one go
    rows = (
        q.join(QuizAttempt, QuizAttempt.id == QuizAttempt.question_id)
        .add_columns(
            QuizAttempt.question.label("question_text"),
            QuizAttempt.type,
            QuizAttempt.correct_option,
            QuizAttempt.correct_answer,
            QuizAttempt.options,
        )
        .all()
    )

    return [
        {
            "question_id": r.QuizAttempt.question_id,
            "user_answer": r.QuizAttempt.user_answer,
            "verdict": r.QuizAttempt.verdict,
            "explanation": r.QuizAttempt.explanation,
            "score": r.QuizAttempt.score,
            "question_text": r.question_text,
            "type": r.type,
            "correct_option": r.correct_option,
            "correct_answer": r.correct_answer,
            "options": r.options,
        }
        for r in rows
    ]

@router.get("/quiz-attempts/{session_id}/list")
def list_attempt_numbers(
    session_id: UUID,
    db: db_dependency,
    attempt: int = Query(None, description="Attempt number (optional)"),    
    current_user: dict = Depends(get_current_user),
):
    nums = (
        db.query(QuizAttempt.attempt_number)
        .filter_by(user_id=current_user.id, session_id=session_id)
        .distinct()
        .order_by(QuizAttempt.attempt_number)
        .all()
    )
    return [n.attempt_number for n in nums]

#DELETE

@router.delete("/study-sessions/{session_id}")
def delete_session(session_id: UUID, db: db_dependency, current_user: dict = Depends(get_current_user)):
    session = db.query(StudySessions).filter_by(id=session_id, user_id=current_user.id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Study session not found")

    # Also delete associated quiz questions
    db.query(QuizQuestion).filter_by(session_id=session_id, user_id=current_user.id).delete()

    db.delete(session)
    db.commit()
    return {"detail": "Session deleted successfully"}

@router.put("/study-rename-session/{session_id}")
def rename_session(session_id: UUID, db: Session = Depends(get_db), user=Depends(get_current_user), data: dict = Body(...)):
    session = db.query(StudySessions).filter_by(id=session_id, user_id=user.id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    session.title = data.get("title", session.title)
    db.commit()
    return {"message": "Renamed successfully"}


# generating docx
# @router.post("/generate-docx")
# async def generate_docx(note: dict):
#     content = note.get("content", "")
#     title = note.get("title", "Notes")

#     # 1️⃣ Convert Markdown → HTML
#     html_content = markdown(content)

#     # 2️⃣ Wrap with title + footer
#     full_html = f"""
#         <h1>{title}</h1>
#         {html_content}
#         <hr/>
#         <p><em>Generated by Learnee • {datetime.now().strftime("%d-%m-%Y")}</em></p>
#     """

#     # 3️⃣ Convert HTML → DOCX
#     file_stream = BytesIO()
#     html2docx(full_html, file_stream)
#     file_stream.seek(0)

#     return Response(
#         content=file_stream.read(),
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         headers={
#             "Content-Disposition": f"attachment; filename={title.replace(' ', '_')}.docx"
#         },
#     )


def add_inline_formatting(paragraph, token):
    if not hasattr(token, "children") or token.children is None:
        paragraph.add_run(token.content if hasattr(token, "content") else "")
        return

    for child in token.children:
        if child.type == "text":
            paragraph.add_run(child.content)

        elif child.type == "strong_open":
            run = paragraph.add_run()
            run.bold = True

        elif child.type == "strong_close":
            pass

        elif child.type == "em_open":
            run = paragraph.add_run()
            run.italic = True

        elif child.type == "em_close":
            pass

        elif child.type == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"

        elif child.type == "softbreak":
            paragraph.add_run("\n")

        elif child.type == "hardbreak":
            paragraph.add_run("\n")


@router.post("/generate-docx")
async def generate_docx(note: dict):
    content = note["content"]
    title = note["title"]

    md = MarkdownIt()
    tokens = md.parse(content)

    document = Document()
    document.add_heading(title, level=1)

    i = 0
    while i < len(tokens):
        t = tokens[i]

        # Headings
        if t.type == "heading_open":
            level = int(t.tag[-1])
            text = tokens[i + 1].content
            document.add_heading(text, level=level)
            i += 3
            continue

        # Paragraph
        if t.type == "paragraph_open":
            p = document.add_paragraph()
            inline = tokens[i + 1]
            add_inline_formatting(p, inline)
            i += 3
            continue

        # Bullet List
        if t.type == "list_item_open" and tokens[i - 1].type == "bullet_list_open":
            p = document.add_paragraph(style="List Bullet")
            inline = tokens[i + 1]
            add_inline_formatting(p, inline)
            i += 3
            continue

        # Numbered List
        if t.type == "list_item_open" and tokens[i - 1].type == "ordered_list_open":
            p = document.add_paragraph(style="List Number")
            inline = tokens[i + 1]
            add_inline_formatting(p, inline)
            i += 3
            continue

        i += 1

    # Footer
    footer = document.add_paragraph(
        f"Generated by Learnee • {datetime.now().strftime('%d-%m-%Y')}"
    )
    footer.italic = True

    # Save DOCX
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)

    return Response(
        content=stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={title.replace(' ', '_')}.docx"
        },
    )
